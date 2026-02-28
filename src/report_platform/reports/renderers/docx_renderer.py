"""DOCX renderer — converts assembled Markdown to a Word document using python-docx.

Handles the Markdown subset used in report chapters:
  - Headings: # H1 through #### H4
  - Pipe-delimited tables: | col1 | col2 |
  - Bold: **text**
  - Italic: *text*
  - Bullet lists: - item or * item
  - Numbered lists: 1. item
  - Horizontal rules: --- (rendered as page breaks)
  - Block quotes: > text
"""

import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Style constants (matching policy_analysis pattern)
NAVY = RGBColor(0, 51, 102)
FONT_NAME = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(24)
FONT_SIZE_H2 = Pt(18)
FONT_SIZE_H3 = Pt(14)
FONT_SIZE_H4 = Pt(12)

# Regex patterns
RE_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
RE_TABLE_SEP = re.compile(r"^\|[\s\-:]+\|$")
RE_BULLET = re.compile(r"^[\-\*]\s+(.+)$")
RE_NUMBERED = re.compile(r"^\d+\.\s+(.+)$")
RE_BLOCKQUOTE = re.compile(r"^>\s?(.*)$")
RE_HR = re.compile(r"^-{3,}$")
RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
RE_BOLD_ITALIC = re.compile(r"\*\*\*(.+?)\*\*\*")


class DocxRenderer:
    """Render Markdown content to DOCX format."""

    def render(self, markdown_content: str, output_path: Path, title: str = "") -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_styles(doc)

        lines = markdown_content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Horizontal rule → page break (between chapters)
            if RE_HR.match(stripped):
                doc.add_page_break()
                i += 1
                continue

            # Heading
            m = RE_HEADING.match(stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                heading = doc.add_heading(level=level)
                self._add_formatted_runs(heading, text)
                i += 1
                continue

            # Table (collect all table rows)
            if RE_TABLE_ROW.match(stripped):
                table_lines = []
                while i < len(lines) and RE_TABLE_ROW.match(lines[i].strip()):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Bullet list
            m = RE_BULLET.match(stripped)
            if m:
                para = doc.add_paragraph(style="List Bullet")
                self._add_formatted_runs(para, m.group(1))
                i += 1
                continue

            # Numbered list
            m = RE_NUMBERED.match(stripped)
            if m:
                para = doc.add_paragraph(style="List Number")
                self._add_formatted_runs(para, m.group(1))
                i += 1
                continue

            # Block quote
            m = RE_BLOCKQUOTE.match(stripped)
            if m:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1.5)
                run = para.add_run(m.group(1))
                run.italic = True
                i += 1
                continue

            # Empty line — skip
            if not stripped:
                i += 1
                continue

            # Regular paragraph
            para = doc.add_paragraph()
            self._add_formatted_runs(para, stripped)
            i += 1

        doc.save(str(output_path))
        logger.info("Wrote DOCX: %s", output_path)
        return output_path

    def _setup_styles(self, doc: Document):
        """Configure document styles — Calibri, navy headings."""
        # Normal body style
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE_BODY

        # Set paragraph spacing
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        # Heading styles
        heading_sizes = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3, 4: FONT_SIZE_H4}
        for level, size in heading_sizes.items():
            hstyle = doc.styles[f"Heading {level}"]
            hstyle.font.color.rgb = NAVY
            hstyle.font.bold = True
            hstyle.font.size = size
            hstyle.font.name = FONT_NAME
            # Add space before headings
            hstyle.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            hstyle.paragraph_format.space_after = Pt(6)

    def _add_formatted_runs(self, paragraph, text: str):
        """Parse inline Markdown formatting (bold, italic) and add runs."""
        # Split text by bold/italic markers and create runs
        parts = self._parse_inline(text)
        for content, bold, italic in parts:
            if not content:
                continue
            run = paragraph.add_run(content)
            run.bold = bold
            run.italic = italic

    def _parse_inline(self, text: str) -> list[tuple[str, bool, bool]]:
        """Parse inline Markdown into (text, bold, italic) tuples."""
        results: list[tuple[str, bool, bool]] = []
        pos = 0

        while pos < len(text):
            # Bold+italic: ***text***
            m = RE_BOLD_ITALIC.search(text, pos)
            m_bold = RE_BOLD.search(text, pos)
            m_italic = RE_ITALIC.search(text, pos)

            # Find the earliest match
            earliest = None
            for candidate in [m, m_bold, m_italic]:
                if candidate and (earliest is None or candidate.start() < earliest.start()):
                    earliest = candidate

            if earliest is None:
                # No more formatting — rest is plain text
                results.append((text[pos:], False, False))
                break

            # Add plain text before the match
            if earliest.start() > pos:
                results.append((text[pos:earliest.start()], False, False))

            if earliest is m:
                results.append((m.group(1), True, True))
            elif earliest is m_bold:
                results.append((m_bold.group(1), True, False))
            else:
                results.append((m_italic.group(1), False, True))

            pos = earliest.end()

        return results if results else [("", False, False)]

    def _add_table(self, doc: Document, table_lines: list[str]):
        """Parse pipe-delimited table lines and add a styled Word table."""
        # Parse rows
        rows: list[list[str]] = []
        for line in table_lines:
            if RE_TABLE_SEP.match(line):
                continue  # Skip separator row
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        # Create table
        n_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply a clean style
        try:
            table.style = doc.styles["Light Grid Accent 1"]
        except KeyError:
            table.style = "Table Grid"

        # Fill cells
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < n_cols:
                    cell = table.cell(i, j)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    self._add_formatted_runs(para, cell_text)

                    # Bold the header row
                    if i == 0:
                        for run in para.runs:
                            run.bold = True

        # Add a blank paragraph after the table for spacing
        doc.add_paragraph()
