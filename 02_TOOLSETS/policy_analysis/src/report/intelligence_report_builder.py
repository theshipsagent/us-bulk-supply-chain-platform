"""
US Flag Fleet Intelligence Brief
Terse. Data-first. Zero fluff.
For maritime industry professionals who don't need basics explained.
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


class IntelligenceReportBuilder:
    """Intelligence brief builder - data tables, vessel names, operators, locations."""

    def __init__(self, config: dict, output_dir: Path):
        self.config = config
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Compact professional styling"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)

        for i in range(1, 4):
            heading = self.doc.styles[f'Heading {i}']
            heading.font.color.rgb = RGBColor(0, 51, 102)
            heading.font.bold = True
            if i == 1:
                heading.font.size = Pt(14)
            elif i == 2:
                heading.font.size = Pt(12)
            else:
                heading.font.size = Pt(11)

    def generate_report(self, vessel_df: pd.DataFrame, analysis: Dict) -> Path:
        """Generate intelligence brief"""
        logger.info("Generating intelligence brief")

        # Prepare key subsets
        msc_df = vessel_df[vessel_df['Vessel Name'].str.contains('USNS', na=False)]
        rrf_df = vessel_df[vessel_df['Vessel Type'].str.contains('Reserve', na=False)]
        commercial_df = vessel_df[~vessel_df['Vessel Name'].str.contains('USNS', na=False) &
                                  ~vessel_df['Vessel Type'].str.contains('Reserve', na=False)]
        norfolk_df = vessel_df[vessel_df['Home Port'] == 'Norfolk']

        # Cover
        self._add_cover()

        # Section 1: Fleet Snapshot
        self._add_fleet_snapshot(vessel_df, analysis, msc_df, rrf_df, commercial_df, norfolk_df)

        # Section 2: MSC Fleet Detail
        self._add_msc_detail(msc_df, vessel_df)

        # Section 3: Norfolk Operations
        self._add_norfolk_ops(norfolk_df)

        # Section 4: Commercial Fleet
        self._add_commercial_fleet(commercial_df)

        # Section 5: Operator Directory
        self._add_operator_directory(vessel_df)

        # Section 6: Business Intelligence
        self._add_business_intel(vessel_df, msc_df, norfolk_df, commercial_df)

        # Appendix: Complete Vessel List
        self._add_vessel_inventory(vessel_df)

        # Save
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.output_dir / f'US_Flag_Intel_Brief_{timestamp}.docx'
        self.doc.save(str(output_path))

        logger.info(f"Intelligence brief generated: {output_path}")
        return output_path

    def _add_cover(self):
        """Minimal cover"""
        title = self.doc.add_heading('US FLAG FLEET', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self.doc.add_heading('Intelligence Brief', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"\n{datetime.now().strftime('%d %B %Y')}\n\n").bold = True
        p.add_run("MSC Fleet | Norfolk Ops | Operator Analysis\n")
        p.add_run("Unclassified / Public Sources")

        self.doc.add_page_break()

    def _add_fleet_snapshot(self, vessel_df, analysis, msc_df, rrf_df, commercial_df, norfolk_df):
        """One-page fleet snapshot. Numbers only."""
        self.doc.add_heading('Fleet Snapshot', level=1)

        # Core numbers
        data = [
            ['Metric', 'Count', '%'],
            ['Total US-Flag Vessels', str(len(vessel_df)), '100%'],
            ['Commercial', str(len(commercial_df)), f"{len(commercial_df)/len(vessel_df)*100:.0f}%"],
            ['MSC (USNS)', str(len(msc_df)), f"{len(msc_df)/len(vessel_df)*100:.0f}%"],
            ['RRF (CAPE-class)', str(len(rrf_df)), f"{len(rrf_df)/len(vessel_df)*100:.0f}%"],
        ]
        self._add_table(data)

        self.doc.add_paragraph()

        # Geographic concentration
        self.doc.add_heading('Homeport Concentration', level=2)
        port_counts = vessel_df['Home Port'].value_counts().head(10)
        port_data = [['Homeport', 'Vessels', 'Primary Operators']]
        for port, count in port_counts.items():
            ops = vessel_df[vessel_df['Home Port'] == port]['Operator'].value_counts().head(2)
            op_str = ', '.join([f"{op} ({ct})" for op, ct in ops.items()])
            port_data.append([port, str(count), op_str])
        self._add_table(port_data)

        self.doc.add_paragraph()

        # Fleet type breakdown
        self.doc.add_heading('Vessel Types', level=2)
        type_counts = vessel_df['Vessel Type'].value_counts()
        type_data = [['Type', 'Count', 'Primary Operators']]
        for vtype, count in type_counts.items():
            ops = vessel_df[vessel_df['Vessel Type'] == vtype]['Operator'].value_counts().head(2)
            op_str = ', '.join(ops.index[:2])
            type_data.append([vtype, str(count), op_str])
        self._add_table(type_data)

        self.doc.add_page_break()

    def _add_msc_detail(self, msc_df, full_df):
        """MSC Fleet - the priority section. Vessel-by-vessel."""
        self.doc.add_heading('MSC Fleet Detail', level=1)

        p = self.doc.add_paragraph()
        p.add_run(f"{len(msc_df)} USNS vessels. ").bold = True
        norfolk_msc = len(msc_df[msc_df['Home Port'] == 'Norfolk'])
        p.add_run(f"{norfolk_msc} at Norfolk. Operated by civilian mariners (CIVMAR) under contract.")

        self.doc.add_paragraph()

        # Group by vessel type/mission area
        type_groups = {
            'Fleet Replenishment Oilers (T-AO)': msc_df[msc_df['Vessel Type'].str.contains('Oiler', na=False)],
            'Fast Combat Support (T-AOE)': msc_df[msc_df['Vessel Type'].str.contains('Fast Combat', na=False)],
            'Large Medium-Speed RoRo (LMSR)': msc_df[msc_df['Vessel Type'].str.contains('Medium-Speed', na=False)],
            'Expeditionary Mobile/Sea Base (ESB/EMB)': msc_df[msc_df['Vessel Type'].str.contains('Expeditionary.*Base', na=False)],
            'Expeditionary Fast Transport (EPF)': msc_df[msc_df['Vessel Type'].str.contains('Fast Transport', na=False)],
            'Hospital Ships (T-AH)': msc_df[msc_df['Vessel Type'].str.contains('Hospital', na=False)],
        }

        for group_name, group_df in type_groups.items():
            if len(group_df) == 0:
                continue

            self.doc.add_heading(group_name, level=2)

            data = [['Vessel', 'Homeport', 'Operator', 'Status', 'Built']]
            for _, v in group_df.sort_values('Vessel Name').iterrows():
                data.append([
                    v['Vessel Name'],
                    v['Home Port'],
                    v['Operator'],
                    v['Status'],
                    str(int(v['Build Year'])) if pd.notna(v['Build Year']) else '-'
                ])
            self._add_table(data)

            # Brief operator note
            ops = group_df['Operator'].value_counts()
            if len(ops) > 0:
                p = self.doc.add_paragraph()
                p.add_run("Operators: ").bold = True
                op_list = [f"{op} ({ct})" for op, ct in ops.items()]
                p.add_run(', '.join(op_list))

            self.doc.add_paragraph()

        # MSC Operator Summary
        self.doc.add_heading('MSC Civilian Contractors', level=2)

        p = self.doc.add_paragraph()
        p.add_run("Ship management contracts by operator:\n")

        msc_ops = msc_df.groupby('Operator').agg({
            'Vessel Name': lambda x: list(x),
            'Home Port': lambda x: list(x.unique())
        })

        for op, row in msc_ops.iterrows():
            p = self.doc.add_paragraph()
            p.add_run(f"{op}: ").bold = True
            p.add_run(f"{len(row['Vessel Name'])} vessels")
            p.add_run(f"\n   Vessels: {', '.join(sorted(row['Vessel Name']))}")
            p.add_run(f"\n   Ports: {', '.join(row['Home Port'])}")

        self.doc.add_page_break()

    def _add_norfolk_ops(self, norfolk_df):
        """Norfolk/Hampton Roads operations"""
        self.doc.add_heading('Norfolk Operations', level=1)

        p = self.doc.add_paragraph()
        p.add_run(f"{len(norfolk_df)} vessels homeported Norfolk. ").bold = True

        msc_nfk = len(norfolk_df[norfolk_df['Vessel Name'].str.contains('USNS', na=False)])
        comm_nfk = len(norfolk_df) - msc_nfk
        p.add_run(f"{msc_nfk} MSC, {comm_nfk} commercial/other.")

        self.doc.add_paragraph()

        # By operator
        self.doc.add_heading('Norfolk by Operator', level=2)
        nfk_ops = norfolk_df.groupby('Operator').agg({
            'Vessel Name': 'count',
            'Vessel Type': lambda x: ', '.join(x.unique())
        }).sort_values('Vessel Name', ascending=False)

        data = [['Operator', 'Vessels', 'Types']]
        for op, row in nfk_ops.iterrows():
            data.append([op, str(row['Vessel Name']), row['Vessel Type']])
        self._add_table(data)

        self.doc.add_paragraph()

        # Complete Norfolk inventory
        self.doc.add_heading('Norfolk Vessel Inventory', level=2)
        data = [['Vessel', 'Type', 'Operator', 'GT', 'Status']]
        for _, v in norfolk_df.sort_values(['Operator', 'Vessel Name']).iterrows():
            data.append([
                v['Vessel Name'],
                v['Vessel Type'],
                v['Operator'],
                f"{int(v['Gross Tonnage']):,}" if pd.notna(v['Gross Tonnage']) else '-',
                v['Status']
            ])
        self._add_table(data)

        self.doc.add_page_break()

    def _add_commercial_fleet(self, commercial_df):
        """Commercial fleet summary - Jones Act, MSP, pure commercial"""
        self.doc.add_heading('Commercial Fleet', level=1)

        p = self.doc.add_paragraph()
        p.add_run(f"{len(commercial_df)} commercial vessels.").bold = True

        self.doc.add_paragraph()

        # Container ships
        containers = commercial_df[commercial_df['Vessel Type'].str.contains('Container', na=False)]
        if len(containers) > 0:
            self.doc.add_heading(f'Container Ships ({len(containers)})', level=2)
            data = [['Vessel', 'Operator', 'Homeport', 'GT', 'Built']]
            for _, v in containers.sort_values(['Operator', 'Vessel Name']).iterrows():
                data.append([
                    v['Vessel Name'],
                    v['Operator'],
                    v['Home Port'],
                    f"{int(v['Gross Tonnage']):,}" if pd.notna(v['Gross Tonnage']) else '-',
                    str(int(v['Build Year'])) if pd.notna(v['Build Year']) else '-'
                ])
            self._add_table(data)

            # Top container operators
            p = self.doc.add_paragraph()
            p.add_run("Top operators: ").bold = True
            cont_ops = containers['Operator'].value_counts().head(5)
            p.add_run(', '.join([f"{op} ({ct})" for op, ct in cont_ops.items()]))
            self.doc.add_paragraph()

        # Tankers
        tankers = commercial_df[commercial_df['Vessel Type'].str.contains('Tanker', na=False)]
        if len(tankers) > 0:
            self.doc.add_heading(f'Tankers ({len(tankers)})', level=2)
            data = [['Vessel', 'Type', 'Operator', 'Homeport', 'DWT']]
            for _, v in tankers.sort_values(['Operator', 'Vessel Name']).iterrows():
                data.append([
                    v['Vessel Name'],
                    v['Vessel Type'],
                    v['Operator'],
                    v['Home Port'],
                    f"{int(v['Deadweight']):,}" if pd.notna(v['Deadweight']) else '-'
                ])
            self._add_table(data)

            p = self.doc.add_paragraph()
            p.add_run("Top operators: ").bold = True
            tank_ops = tankers['Operator'].value_counts().head(5)
            p.add_run(', '.join([f"{op} ({ct})" for op, ct in tank_ops.items()]))
            self.doc.add_paragraph()

        # RoRo/ConRo
        roro = commercial_df[commercial_df['Vessel Type'].str.contains('RoRo|Roll', case=False, na=False)]
        if len(roro) > 0:
            self.doc.add_heading(f'RoRo/ConRo ({len(roro)})', level=2)
            data = [['Vessel', 'Operator', 'Homeport', 'GT']]
            for _, v in roro.sort_values(['Operator', 'Vessel Name']).iterrows():
                data.append([
                    v['Vessel Name'],
                    v['Operator'],
                    v['Home Port'],
                    f"{int(v['Gross Tonnage']):,}" if pd.notna(v['Gross Tonnage']) else '-'
                ])
            self._add_table(data)
            self.doc.add_paragraph()

        # Bulk carriers
        bulk = commercial_df[commercial_df['Vessel Type'].str.contains('Bulk|Freighter', na=False)]
        if len(bulk) > 0:
            self.doc.add_heading(f'Bulk/Lake Freighters ({len(bulk)})', level=2)
            data = [['Vessel', 'Operator', 'Homeport', 'DWT']]
            for _, v in bulk.sort_values(['Operator', 'Vessel Name']).iterrows():
                data.append([
                    v['Vessel Name'],
                    v['Operator'],
                    v['Home Port'],
                    f"{int(v['Deadweight']):,}" if pd.notna(v['Deadweight']) else '-'
                ])
            self._add_table(data)
            self.doc.add_paragraph()

        self.doc.add_page_break()

    def _add_operator_directory(self, vessel_df):
        """Complete operator directory"""
        self.doc.add_heading('Operator Directory', level=1)

        ops = vessel_df.groupby('Operator').agg({
            'Vessel Name': list,
            'Vessel Type': lambda x: list(x.unique()),
            'Home Port': lambda x: list(x.unique()),
            'Gross Tonnage': 'sum'
        }).sort_values('Gross Tonnage', ascending=False)

        for op, row in ops.iterrows():
            p = self.doc.add_paragraph()
            run = p.add_run(f"{op}")
            run.bold = True
            run.font.size = Pt(11)

            p = self.doc.add_paragraph()
            p.add_run(f"Vessels ({len(row['Vessel Name'])}): ").bold = True
            p.add_run(', '.join(sorted(row['Vessel Name'])))

            p = self.doc.add_paragraph()
            p.add_run("Types: ").bold = True
            p.add_run(', '.join(row['Vessel Type']))

            p = self.doc.add_paragraph()
            p.add_run("Homeports: ").bold = True
            p.add_run(', '.join(row['Home Port']))

            p = self.doc.add_paragraph()
            p.add_run("Total GT: ").bold = True
            gt = row['Gross Tonnage']
            p.add_run(f"{int(gt):,}" if pd.notna(gt) else '-')

            self.doc.add_paragraph()  # Spacer

        self.doc.add_page_break()

    def _add_business_intel(self, vessel_df, msc_df, norfolk_df, commercial_df):
        """Business Intelligence - targeting opportunities"""
        self.doc.add_heading('Business Intelligence', level=1)

        # Norfolk operators (BD targets)
        self.doc.add_heading('Norfolk Market', level=2)

        p = self.doc.add_paragraph()
        p.add_run("BD targets by vessel concentration:")

        nfk_ops = norfolk_df['Operator'].value_counts()
        data = [['Operator', 'Norfolk Vessels', 'Fleet Total', 'Norfolk %']]
        for op, nfk_count in nfk_ops.items():
            total = len(vessel_df[vessel_df['Operator'] == op])
            pct = nfk_count / total * 100 if total > 0 else 0
            data.append([op, str(nfk_count), str(total), f"{pct:.0f}%"])
        self._add_table(data)

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.add_run("Key insight: ").bold = True
        top_nfk = nfk_ops.head(3)
        p.add_run(f"{', '.join(top_nfk.index[:2])} dominate Norfolk ops. Combined {top_nfk.sum()} vessels.")

        self.doc.add_paragraph()

        # MSC contractor landscape
        self.doc.add_heading('MSC Contract Landscape', level=2)

        msc_contractors = msc_df['Operator'].value_counts()
        data = [['Contractor', 'MSC Vessels', 'Vessel Types', 'Primary Homeport']]
        for op, count in msc_contractors.items():
            op_vessels = msc_df[msc_df['Operator'] == op]
            types = ', '.join(op_vessels['Vessel Type'].unique()[:2])
            ports = op_vessels['Home Port'].value_counts()
            primary_port = ports.index[0] if len(ports) > 0 else '-'
            data.append([op, str(count), types, primary_port])
        self._add_table(data)

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.add_run("MSC ship management concentrated: ").bold = True
        top2 = msc_contractors.head(2)
        p.add_run(f"{top2.index[0]} ({top2.iloc[0]}), {top2.index[1]} ({top2.iloc[1]}) hold {top2.sum()}/{len(msc_df)} MSC vessels.")

        self.doc.add_paragraph()

        # Port service opportunities
        self.doc.add_heading('Port Service Opportunities', level=2)

        p = self.doc.add_paragraph()
        p.add_run("Ports by vessel concentration (agency/terminal targets):\n")

        port_data = vessel_df.groupby('Home Port').agg({
            'Vessel Name': 'count',
            'Gross Tonnage': 'sum',
            'Operator': lambda x: len(x.unique())
        }).sort_values('Vessel Name', ascending=False)

        data = [['Port', 'Vessels', 'Operators', 'Total GT']]
        for port, row in port_data.head(10).iterrows():
            data.append([
                port,
                str(row['Vessel Name']),
                str(row['Operator']),
                f"{int(row['Gross Tonnage']):,}" if pd.notna(row['Gross Tonnage']) else '-'
            ])
        self._add_table(data)

        self.doc.add_paragraph()

        # Jones Act market
        self.doc.add_heading('Jones Act Fleet', level=2)

        jones_types = ['Container Ship', 'Product Tanker', 'Crude Oil Tanker', 'RoRo Cargo']
        jones_df = commercial_df[commercial_df['Vessel Type'].isin(jones_types)]

        p = self.doc.add_paragraph()
        p.add_run(f"{len(jones_df)} Jones Act-eligible vessels. ").bold = True

        ja_ops = jones_df['Operator'].value_counts()
        p.add_run(f"Top operators: {', '.join([f'{op} ({ct})' for op, ct in ja_ops.head(5).items()])}")

        self.doc.add_paragraph()

        # Market assessment
        self.doc.add_heading('Market Assessment', level=2)

        bullets = [
            f"Norfolk: {len(norfolk_df)} vessels, dominated by MSC ops. High service demand.",
            f"West Coast: Container hub - Matson, Crowley, APL. Hawaii/Alaska trades.",
            f"Gulf Coast: Tanker concentration. OSG, Liberty Maritime, American Shipping.",
            f"Great Lakes: Seasonal bulk ops. Interlake, Great Lakes Fleet.",
            f"MSC contracts: Maersk Line Ltd largest. Crowley holds ESB/EMB portfolio."
        ]

        for bullet in bullets:
            p = self.doc.add_paragraph()
            p.add_run("• " + bullet)

        self.doc.add_page_break()

    def _add_vessel_inventory(self, vessel_df):
        """Complete vessel inventory"""
        self.doc.add_heading('Appendix: Complete Vessel Inventory', level=1)

        p = self.doc.add_paragraph()
        p.add_run(f"{len(vessel_df)} vessels. Sorted by operator.")

        self.doc.add_paragraph()

        data = [['Vessel', 'Type', 'Operator', 'Port', 'GT', 'DWT', 'Built']]
        for _, v in vessel_df.sort_values(['Operator', 'Vessel Name']).iterrows():
            data.append([
                v['Vessel Name'],
                v['Vessel Type'][:20] if len(str(v['Vessel Type'])) > 20 else v['Vessel Type'],
                v['Operator'][:20] if len(str(v['Operator'])) > 20 else v['Operator'],
                v['Home Port'],
                f"{int(v['Gross Tonnage']):,}" if pd.notna(v['Gross Tonnage']) else '-',
                f"{int(v['Deadweight']):,}" if pd.notna(v['Deadweight']) else '-',
                str(int(v['Build Year'])) if pd.notna(v['Build Year']) else '-'
            ])
        self._add_table(data, small=True)

    def _add_table(self, data: List[List[str]], small: bool = False):
        """Add compact table"""
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_value) if cell_value is not None else '-'
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(8) if small else Pt(9)
                        if i == 0:
                            run.font.bold = True

        self.doc.add_paragraph()
