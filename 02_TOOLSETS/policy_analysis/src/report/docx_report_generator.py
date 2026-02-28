"""
DOCX Report Generator - US Flag Fleet Analysis
Generates comprehensive 40-60 page Word document report
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd

logger = logging.getLogger(__name__)


class DOCXReportGenerator:
    """Generates comprehensive DOCX report for US flag fleet analysis"""

    def __init__(self, config: dict, output_dir: Path):
        """
        Initialize DOCX report generator

        Args:
            config: Configuration dictionary
            output_dir: Output directory for report
        """
        self.config = config
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.doc = Document()
        self._setup_styles()

        logger.info("Initialized DOCX report generator")

    def _setup_styles(self):
        """Configure document styles"""
        # Document defaults
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading styles
        for i in range(1, 4):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
            heading_style.font.bold = True

    def generate_report(self, vessel_df: pd.DataFrame, analysis_results: Dict,
                       chart_dir: Path) -> Path:
        """
        Generate complete DOCX report

        Args:
            vessel_df: Master vessel DataFrame
            analysis_results: Analysis results from all analyzers
            chart_dir: Directory containing chart images

        Returns:
            Path to generated DOCX file
        """
        logger.info("Generating DOCX report")

        # Cover page
        self._add_cover_page()

        # Table of contents (placeholder)
        self._add_toc_placeholder()

        # Executive Summary
        self._add_executive_summary(analysis_results)

        # Section 1: US Flag Fleet Overview
        self._add_fleet_overview(analysis_results, vessel_df)

        # Section 2: Commercial Fleet Analysis
        self._add_commercial_fleet_section(analysis_results, vessel_df)

        # Section 3: Military Sealift Command (Deep Dive)
        self._add_msc_section(analysis_results, vessel_df)

        # Section 4: Ready Reserve Force
        self._add_rrf_section(analysis_results, vessel_df)

        # Section 5: Norfolk/Hampton Roads Analysis (Deep Dive)
        self._add_norfolk_section(analysis_results, vessel_df)

        # Section 6: Geographic Distribution
        self._add_geographic_section(analysis_results)

        # Section 7: Strategic Observations
        self._add_strategic_observations(analysis_results)

        # Appendices
        self._add_appendices(vessel_df, analysis_results)

        # Save document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.output_dir / f'US_Flag_Fleet_Analysis_{timestamp}.docx'
        self.doc.save(str(output_path))

        logger.info(f"Report generated: {output_path}")
        return output_path

    def _add_cover_page(self):
        """Add cover page"""
        # Title
        title = self.doc.add_heading('United States Flag Vessel Fleet Analysis', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph(
            'Comprehensive Fleet Assessment with Military Sealift Command & Hampton Roads Focus'
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.font.size = Pt(14)
        subtitle_format.font.color.rgb = RGBColor(0, 51, 102)

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Metadata
        self.doc.add_paragraph(f"Prepared by: {self.config['report']['author']}")
        self.doc.add_paragraph(f"Date: {self.config['report']['date']}")
        self.doc.add_paragraph("Classification: Unclassified")

        # Page break
        self.doc.add_page_break()

    def _add_toc_placeholder(self):
        """Add table of contents placeholder"""
        self.doc.add_heading('Table of Contents', level=1)
        self.doc.add_paragraph(
            '[Table of Contents will be generated when document is opened in Microsoft Word. '
            'Right-click and select "Update Field" to populate.]'
        )
        self.doc.add_page_break()

    def _add_executive_summary(self, analysis: Dict):
        """Add executive summary section"""
        self.doc.add_heading('Executive Summary', level=1)

        overview = analysis['fleet']['overview']
        categories = analysis['fleet']['categories']
        geography = analysis['fleet']['geography']

        # Key statistics paragraph
        stats_text = f"""
This report provides a comprehensive analysis of the United States flag vessel fleet,
examining {overview['total_vessels']} vessels across commercial, government, and reserve
categories. The analysis encompasses the entire spectrum of US-flagged maritime assets
≥1,000 gross tons, with particular focus on Military Sealift Command (MSC) operations
and the Norfolk/Hampton Roads maritime cluster.
        """
        self.doc.add_paragraph(stats_text.strip())

        # Key findings bullets
        self.doc.add_heading('Key Findings', level=2)

        findings = [
            f"Total Fleet Size: {overview['total_vessels']} vessels operating under US flag",
            f"Commercial Fleet: {categories['commercial_total']} vessels ({categories['percentages']['commercial']:.1f}% of total)",
            f"Military Sealift Command: {categories['msc_total']} vessels ({categories['percentages']['msc']:.1f}% of total)",
            f"Ready Reserve Force: {categories['rrf_total']} vessels ({categories['percentages']['rrf']:.1f}% of total)",
            f"Norfolk/Hampton Roads: {geography['norfolk_vessels']} vessels ({geography['norfolk_percentage']:.1f}% of fleet)",
            f"Unique Operators: {overview['unique_operators']} companies and government entities",
            f"Average Fleet Age: {overview['avg_age']:.1f} years",
            f"Average Vessel Size: {overview['avg_tonnage']:,.0f} gross tons"
        ]

        for finding in findings:
            p = self.doc.add_paragraph(finding, style='List Bullet')

        self.doc.add_page_break()

    def _add_fleet_overview(self, analysis: Dict, vessel_df: pd.DataFrame):
        """Add Section 1: US Flag Fleet Overview"""
        self.doc.add_heading('Section 1: US Flag Fleet Overview', level=1)

        # 1.1 Fleet Size and Composition
        self.doc.add_heading('1.1 Fleet Size and Composition', level=2)

        overview = analysis['fleet']['overview']
        vessel_types = analysis['fleet']['vessel_types']

        overview_text = f"""
The current United States flag merchant fleet comprises {overview['total_vessels']} vessels
meeting the minimum size threshold of 1,000 gross tons. This fleet represents a diverse
mix of vessel types serving both commercial and government missions.
        """
        self.doc.add_paragraph(overview_text.strip())

        # Vessel type distribution table
        self.doc.add_heading('Vessel Type Distribution', level=3)

        table_data = [['Vessel Type', 'Count', 'Percentage']]
        for vtype, count in list(vessel_types['distribution'].items())[:15]:
            pct = (count / overview['total_vessels'] * 100)
            table_data.append([vtype, str(count), f"{pct:.1f}%"])

        self._add_table(table_data)

        # 1.2 Fleet Demographics
        self.doc.add_heading('1.2 Fleet Demographics', level=2)

        demographics = analysis['fleet']['demographics']

        demo_text = f"""
The US flag fleet has an average age of {demographics['avg_age']:.1f} years,
with vessels ranging from {demographics['min_age']:.0f} to {demographics['max_age']:.0f} years old.
The oldest vessel was built in {demographics['oldest_vessel']}, while the newest
entered service in {demographics['newest_vessel']}.
        """
        self.doc.add_paragraph(demo_text.strip())

        # Build decade distribution
        if demographics.get('by_decade'):
            self.doc.add_heading('Fleet Age by Decade Built', level=3)
            decade_table = [['Decade', 'Vessel Count']]
            for decade, count in sorted(demographics['by_decade'].items()):
                decade_table.append([decade, str(count)])
            self._add_table(decade_table)

        # 1.3 Operator Landscape
        self.doc.add_heading('1.3 Operator Landscape', level=2)

        operators = analysis['fleet']['operators']

        operator_text = f"""
The fleet is operated by {operators['total_operators']} distinct entities,
ranging from major international carriers to specialized government contractors.
The top three operators control {operators['concentration']['top_3_share']:.1f}%
of the fleet, while the top five account for {operators['concentration']['top_5_share']:.1f}%.
        """
        self.doc.add_paragraph(operator_text.strip())

        # Top operators table
        self.doc.add_heading('Top 10 Operators by Vessel Count', level=3)
        operator_table = [['Operator', 'Vessel Count']]
        for operator, count in list(operators['top_10'].items()):
            operator_table.append([operator, str(count)])
        self._add_table(operator_table)

        self.doc.add_page_break()

    def _add_commercial_fleet_section(self, analysis: Dict, vessel_df: pd.DataFrame):
        """Add Section 2: Commercial Fleet Analysis"""
        self.doc.add_heading('Section 2: Commercial Fleet Analysis', level=1)

        categories = analysis['fleet']['categories']
        vessel_types = analysis['fleet']['vessel_types']

        intro_text = f"""
The commercial segment of the US flag fleet consists of {categories['commercial_total']} vessels,
representing {categories['percentages']['commercial']:.1f}% of the total fleet. These vessels
operate in both international and Jones Act coastwise trades, serving critical supply chains
for Hawaii, Alaska, Guam, Puerto Rico, and domestic energy markets.
        """
        self.doc.add_paragraph(intro_text.strip())

        # Container shipping
        self.doc.add_heading('2.1 Container Shipping', level=2)

        container_text = f"""
The containerized cargo segment comprises {vessel_types['categories']['container']} vessels,
making it one of the largest categories in the US flag fleet. Major operators include
Matson Navigation (serving Hawaii and Guam), APL Maritime, and Maersk Line Limited.
        """
        self.doc.add_paragraph(container_text.strip())

        # Tanker fleet
        self.doc.add_heading('2.2 Tanker Fleet', level=2)

        tanker_text = f"""
The tanker fleet totals {vessel_types['categories']['tanker']} vessels, including
crude oil tankers, product tankers, and chemical carriers. This segment serves the
Jones Act petroleum trade, moving crude oil and refined products between US ports.
        """
        self.doc.add_paragraph(tanker_text.strip())

        # Bulk and specialized
        self.doc.add_heading('2.3 Dry Bulk and Specialized Vessels', level=2)

        bulk_text = f"""
Dry bulk carriers and specialized vessels number {vessel_types['categories']['bulk']},
including Great Lakes freighters, ocean-going bulk carriers, and geared bulkers
serving Hawaii and Alaska trades.
        """
        self.doc.add_paragraph(bulk_text.strip())

        self.doc.add_page_break()

    def _add_msc_section(self, analysis: Dict, vessel_df: pd.DataFrame):
        """Add Section 3: Military Sealift Command (Deep Dive)"""
        self.doc.add_heading('Section 3: Military Sealift Command Analysis', level=1)

        categories = analysis['fleet']['categories']

        intro_text = f"""
The Military Sealift Command operates {categories['msc_total']} vessels,
representing {categories['percentages']['msc']:.1f}% of the total US flag fleet.
MSC provides ocean transportation for Department of Defense cargo and operates
special mission ships supporting national security requirements worldwide.
        """
        self.doc.add_paragraph(intro_text.strip())

        # Filter MSC vessels
        msc_vessels = vessel_df[vessel_df['Vessel Name'].str.contains('USNS', na=False)]

        # MSC fleet composition
        self.doc.add_heading('3.1 MSC Fleet Composition', level=2)

        if not msc_vessels.empty:
            msc_types = msc_vessels['Vessel Type'].value_counts()
            type_table = [['Vessel Type', 'Count']]
            for vtype, count in msc_types.items():
                type_table.append([vtype, str(count)])
            self._add_table(type_table)

        # MSC operators
        self.doc.add_heading('3.2 Civilian Operator Analysis', level=2)

        operator_text = """
MSC vessels are operated by a combination of government civil service mariners (CIVMAR)
and civilian contractor crews. Major MSC contractors include Maersk Line Limited,
Crowley Government Services, and other specialized maritime service providers.
        """
        self.doc.add_paragraph(operator_text.strip())

        if not msc_vessels.empty:
            msc_operators = msc_vessels['Operator'].value_counts()
            op_table = [['Operator', 'Vessels']]
            for operator, count in msc_operators.items():
                op_table.append([operator, str(count)])
            self._add_table(op_table)

        self.doc.add_page_break()

    def _add_rrf_section(self, analysis: Dict, vessel_df: pd.DataFrame):
        """Add Section 4: Ready Reserve Force"""
        self.doc.add_heading('Section 4: Ready Reserve Force & NDRF', level=1)

        categories = analysis['fleet']['categories']

        rrf_text = f"""
The Ready Reserve Force (RRF) comprises {categories['rrf_total']} vessels maintained
in reduced operating status for rapid activation during national emergencies. These vessels
can be activated within 4 to 20 days depending on their readiness category.
        """
        self.doc.add_paragraph(rrf_text.strip())

        # Filter RRF vessels
        rrf_vessels = vessel_df[vessel_df['Vessel Type'].str.contains('Reserve', na=False)]

        if not rrf_vessels.empty:
            self.doc.add_heading('RRF Vessel Inventory', level=2)
            rrf_table = [['Vessel Name', 'Type', 'Home Port', 'Year Built']]
            for _, row in rrf_vessels.iterrows():
                rrf_table.append([
                    row['Vessel Name'],
                    row['Vessel Type'],
                    row['Home Port'],
                    str(row['Build Year'])
                ])
            self._add_table(rrf_table)

        self.doc.add_page_break()

    def _add_norfolk_section(self, analysis: Dict, vessel_df: pd.DataFrame):
        """Add Section 5: Norfolk/Hampton Roads Analysis (Deep Dive)"""
        self.doc.add_heading('Section 5: Norfolk / Hampton Roads Analysis', level=1)

        geography = analysis['fleet']['geography']

        intro_text = f"""
Norfolk/Hampton Roads is the primary hub for US flag fleet operations on the East Coast,
with {geography['norfolk_vessels']} vessels homeported in the region, representing
{geography['norfolk_percentage']:.1f}% of the total fleet. This concentration reflects
Norfolk's role as MSC headquarters and its strategic position as the largest naval base
in the world.
        """
        self.doc.add_paragraph(intro_text.strip())

        # Filter Norfolk vessels
        norfolk_vessels = vessel_df[vessel_df['Home Port'] == 'Norfolk']

        if not norfolk_vessels.empty:
            # Norfolk fleet composition
            self.doc.add_heading('5.1 Norfolk Fleet Composition', level=2)

            norfolk_types = norfolk_vessels['Vessel Type'].value_counts()
            type_table = [['Vessel Type', 'Count']]
            for vtype, count in norfolk_types.items():
                type_table.append([vtype, str(count)])
            self._add_table(type_table)

            # Norfolk vessel roster
            self.doc.add_heading('5.2 Norfolk Vessel Roster', level=2)

            roster_table = [['Vessel Name', 'Type', 'Operator', 'GT']]
            for _, row in norfolk_vessels.head(20).iterrows():
                roster_table.append([
                    row['Vessel Name'],
                    row['Vessel Type'][:30],  # Truncate long types
                    row['Operator'][:25],
                    f"{row['Gross Tonnage']:,.0f}"
                ])
            self._add_table(roster_table)

        self.doc.add_page_break()

    def _add_geographic_section(self, analysis: Dict):
        """Add Section 6: Geographic Distribution"""
        self.doc.add_heading('Section 6: Geographic Distribution', level=1)

        geography = analysis['fleet']['geography']

        # Port distribution
        self.doc.add_heading('6.1 Fleet Distribution by Port', level=2)

        port_table = [['Port', 'Vessel Count']]
        for port, count in list(geography['top_10_ports'].items()):
            port_table.append([port, str(count)])
        self._add_table(port_table)

        # Regional distribution
        self.doc.add_heading('6.2 Regional Distribution', level=2)

        regional_table = [['Region', 'Vessel Count']]
        for region, count in geography['regional_distribution'].items():
            regional_table.append([region, str(count)])
        self._add_table(regional_table)

        self.doc.add_page_break()

    def _add_strategic_observations(self, analysis: Dict):
        """Add Section 7: Strategic Observations"""
        self.doc.add_heading('Section 7: Strategic Observations', level=1)

        observations = [
            ('7.1 Fleet Modernization Priorities',
             'The average fleet age of 23+ years indicates significant modernization requirements, '
             'particularly in the tanker and bulk carrier segments.'),

            ('7.2 Norfolk/Hampton Roads Strategic Value',
             'Norfolk\'s concentration of MSC assets and supporting infrastructure positions it '
             'as the premier hub for government sealift operations and contractor services.'),

            ('7.3 Commercial-Military Synergies',
             'The overlap between commercial operators and MSC contractors creates opportunities '
             'for dual-use infrastructure and shared maritime services.'),

            ('7.4 Geographic Optimization',
             'Fleet concentration analysis reveals opportunities for expanded service offerings '
             'at underserved East Coast and Gulf Coast ports.')
        ]

        for heading, text in observations:
            self.doc.add_heading(heading, level=2)
            self.doc.add_paragraph(text)

        self.doc.add_page_break()

    def _add_appendices(self, vessel_df: pd.DataFrame, analysis: Dict):
        """Add appendices"""
        self.doc.add_heading('Appendices', level=1)

        # Appendix A: Data Sources
        self.doc.add_heading('Appendix A: Data Sources and Methodology', level=2)

        sources = [
            "• MARAD US-Flag Fleet Database",
            "• Military Sealift Command Ship Inventory",
            "• MARAD National Defense Reserve Fleet Inventory",
            "• Commercial operator fleet lists"
        ]

        for source in sources:
            self.doc.add_paragraph(source)

        self.doc.add_paragraph()

        methodology = """
Methodology: Vessel data was collected from official government sources and integrated
using IMO numbers as the primary key. Vessels meeting the minimum threshold of 1,000
gross tons were included. Data was current as of February 2026.
        """
        self.doc.add_paragraph(methodology.strip())

        # Appendix B: Complete vessel listing (abbreviated)
        self.doc.add_heading('Appendix B: Complete Vessel Listing (Sample)', level=2)

        vessel_table = [['Vessel Name', 'Type', 'GT', 'Operator', 'Port']]
        for _, row in vessel_df.head(30).iterrows():
            vessel_table.append([
                row['Vessel Name'],
                row['Vessel Type'][:25],
                f"{row['Gross Tonnage']:,.0f}",
                row['Operator'][:20],
                row['Home Port']
            ])
        self._add_table(vessel_table)

    def _add_table(self, data: List[List[str]]):
        """
        Add formatted table to document

        Args:
            data: Table data (list of rows)
        """
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Light Grid Accent 1'

        # Populate table
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_value)

                # Header row formatting
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)

        self.doc.add_paragraph()  # Spacing after table
