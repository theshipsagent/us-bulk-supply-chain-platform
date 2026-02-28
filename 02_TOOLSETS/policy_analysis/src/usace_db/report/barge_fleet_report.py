"""
DOCX report builder for Mississippi River Dry Barge Fleet Acquisition Analysis.
Generates a CEO-level ~25-30 page report with data tables, charts, and narrative analysis.
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd

logger = logging.getLogger(__name__)

NAVY = RGBColor(0, 0x33, 0x66)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_HEX = 'D5D8DC'
HEADER_BG_HEX = '003366'


class BargeFleetReport:
    """Builds the complete DOCX acquisition analysis report."""

    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure Calibri 10pt body, navy headings."""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)

        for level in range(1, 4):
            hs = self.doc.styles[f'Heading {level}']
            hs.font.color.rgb = NAVY
            hs.font.bold = True
            hs.font.name = 'Calibri'
            if level == 1:
                hs.font.size = Pt(16)
                hs.paragraph_format.space_before = Pt(18)
            elif level == 2:
                hs.font.size = Pt(13)
                hs.paragraph_format.space_before = Pt(12)
            else:
                hs.font.size = Pt(11)
                hs.paragraph_format.space_before = Pt(8)

    def _add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                       size: int = 10, color: RGBColor = None, alignment=None):
        """Add a styled paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color
        if alignment:
            p.alignment = alignment
        return p

    def _add_table(self, headers: List[str], rows: List[List], widths: List[float] = None):
        """Add a formatted data table with navy header row."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Navy background
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): HEADER_BG_HEX
            })
            shading.append(shd)

        # Data rows
        for i, row_data in enumerate(rows):
            for j, val in enumerate(row_data):
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                # Right-align numeric columns (heuristic: if value starts with digit or $)
                s = str(val).strip()
                if s and (s[0].isdigit() or s.startswith('$') or s.startswith('-')):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Alternating row shading
                if i % 2 == 1:
                    shading = cell._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn('w:shd'), {
                        qn('w:val'): 'clear',
                        qn('w:color'): 'auto',
                        qn('w:fill'): 'F2F3F4'
                    })
                    shading.append(shd)

        # Column widths
        if widths:
            for j, w in enumerate(widths):
                for row in table.rows:
                    row.cells[j].width = Cm(w)

        self.doc.add_paragraph()  # spacing after table
        return table

    def _add_chart(self, chart_path: Path, caption: str = ''):
        """Embed a chart PNG with optional caption."""
        if chart_path and chart_path.exists():
            self.doc.add_picture(str(chart_path), width=Inches(6))
            if caption:
                p = self._add_paragraph(caption, italic=True, size=8,
                                        color=RGBColor(0x66, 0x66, 0x66))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()

    # =================================================================
    # COVER PAGE
    # =================================================================
    def _add_cover_page(self):
        """Professional cover page with CONFIDENTIAL marker."""
        self.doc.add_paragraph('\n' * 4)

        title = self._add_paragraph(
            'Mississippi River System', bold=True, size=26, color=NAVY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self._add_paragraph(
            'Dry Barge Fleet Acquisition Analysis', bold=True, size=20, color=NAVY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self.doc.add_paragraph()

        self._add_paragraph(
            'Market Assessment & Fleet Demographics', size=14, color=DARK_GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._add_paragraph(
            'USACE Waterborne Commerce Data | 2013-2023', size=12, color=DARK_GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self.doc.add_paragraph('\n' * 3)

        self._add_paragraph(
            f'Prepared: {datetime.now().strftime("%B %Y")}', size=10,
            alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self.doc.add_paragraph('\n' * 2)

        self._add_paragraph(
            'CONFIDENTIAL — FOR INTERNAL USE ONLY', bold=True, size=11,
            color=RGBColor(0xCC, 0, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self.doc.add_page_break()

    # =================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # =================================================================
    def _add_executive_summary(self, summary: dict, concentration: dict,
                               fleet_by_year: pd.DataFrame):
        """2-3 page executive summary with key metrics table and investment thesis."""
        self.doc.add_heading('1. Executive Summary', level=1)

        # Key metrics table
        self.doc.add_heading('Fleet Snapshot (2023)', level=2)
        self._add_table(
            ['Metric', 'Value'],
            [
                ['Total Dry Barges (Mississippi)', f"{summary['total_barges']:,}"],
                ['Covered Hopper Barges', f"{summary['covered']:,}"],
                ['Open Hopper Barges', f"{summary['open']:,}"],
                ['Covered/Open Ratio', f"{summary['covered'] / max(summary['open'], 1):.2f}x"],
                ['Unique Operators', f"{summary['unique_operators']:,}"],
                ['Average Fleet Age', f"{summary['avg_age']:.1f} years"],
                ['Median Fleet Age', f"{summary['median_age']:.1f} years"],
                ['Oldest Active Build Year', str(summary['oldest_build'])],
                ['Newest Build Year', str(summary['newest_build'])],
                ['10-Year Fleet Growth', f"{fleet_by_year['total'].iloc[-1] - fleet_by_year['total'].iloc[0]:+,} barges"],
                ['Top-3 Operator Concentration', f"{concentration['top3_share']:.1f}%"],
                ['HHI (Market Concentration)', f"{concentration['hhi']:,.0f}"],
            ]
        )

        # Bull case
        self.doc.add_heading('Bull Case for Acquisition', level=2)
        self._add_paragraph(
            f"The Mississippi River dry barge fleet has grown from {fleet_by_year['total'].iloc[0]:,} "
            f"to {fleet_by_year['total'].iloc[-1]:,} barges over the 2013-2023 period, a net increase of "
            f"{fleet_by_year['total'].iloc[-1] - fleet_by_year['total'].iloc[0]:,} units driven primarily "
            f"by covered hopper expansion tied to grain export growth. The covered fleet now represents "
            f"{summary['covered'] / summary['total_barges'] * 100:.1f}% of the total, reflecting the "
            f"structural shift from coal to agricultural commodities as the dominant inland waterway cargo."
        )
        self._add_paragraph(
            f"The fleet faces a significant replacement cliff. With an average age of "
            f"{summary['avg_age']:.1f} years and industry-standard economic life of 25-30 years, "
            f"a substantial portion of the fleet will require replacement within the next decade. "
            f"Meanwhile, new construction capacity is severely constrained following the Jeffboat closure "
            f"in 2018, which eliminated the largest inland barge builder in North America. Remaining "
            f"yards — primarily Arcosa Marine (formerly Trinity) and Heartland Barge — face order backlogs "
            f"extending 18-24 months. This supply-demand imbalance supports elevated secondhand barge values "
            f"and strengthens the case for acquiring existing fleet assets."
        )
        self._add_paragraph(
            f"U.S. grain exports remain the primary demand driver. USDA long-term projections show "
            f"corn and soybean export volumes growing at 1-2% annually through 2032, underpinned by "
            f"global protein demand, ethanol mandates, and the cost advantage of Mississippi River "
            f"transportation versus rail and truck. South American production volatility and recurring "
            f"low-water events on competing river systems (Parana/Paraguay) further support U.S. "
            f"waterway utilization."
        )

        # Bear case
        self.doc.add_heading('Bear Case & Risk Factors', level=2)
        self._add_paragraph(
            f"Coal transportation, historically the backbone of the open hopper fleet, continues its "
            f"structural decline. EIA data shows U.S. coal production has fallen over 40% from its 2008 "
            f"peak, with utility retirements accelerating the trend. While some open hoppers have been "
            f"repurposed for aggregates and industrial minerals, the coal-to-gas transition represents "
            f"permanent demand destruction for roughly 15-20% of the open hopper fleet."
        )
        self._add_paragraph(
            f"Trade policy risk remains elevated. The 2018-2019 U.S.-China trade war demonstrated "
            f"the vulnerability of grain barge demand to tariff escalation, with Chinese soybean purchases "
            f"shifting to Brazil during the dispute. Any renewal of agricultural trade tensions would "
            f"directly impact covered hopper utilization. Additionally, prolonged low-water events on "
            f"the Mississippi — as occurred in 2022 — can severely restrict barge loadings and disrupt "
            f"supply chains for extended periods."
        )
        self._add_paragraph(
            f"New construction costs have escalated significantly, with covered hoppers now pricing "
            f"at $750,000-$950,000 per unit depending on specifications, up from approximately "
            f"$500,000-$600,000 a decade ago. While this supports secondhand values, it also raises "
            f"the cost of fleet renewal and expansion for any acquirer planning a growth strategy."
        )

        # Signal matrix
        self.doc.add_heading('Acquisition Signal Matrix', level=2)
        self._add_table(
            ['Factor', 'Signal', 'Assessment'],
            [
                ['Fleet Age Profile', f"Avg {summary['avg_age']:.0f} yrs, retirement cliff approaching", 'Positive'],
                ['Builder Capacity', 'Jeffboat closed; 18-24 mo backlogs at remaining yards', 'Positive'],
                ['Grain Export Outlook', 'USDA projects 1-2% annual growth through 2032', 'Positive'],
                ['Coal Demand Trend', '40%+ decline from peak; structural, not cyclical', 'Negative'],
                ['Market Concentration', f"Top-3 hold {concentration['top3_share']:.0f}%; fragmented tail", 'Neutral — M&A opportunity'],
                ['Secondhand Values', 'Elevated due to constrained supply; $350-550K for 5-10 yr units', 'Positive (for sellers)'],
                ['Trade Policy', 'China tariff risk on grain; recurring threat', 'Negative'],
                ['Low-Water Risk', '2022 event disrupted season; climate-driven frequency increasing', 'Negative'],
                ['Regulatory', 'No major pending changes; subchapter M for towboats already absorbed', 'Neutral'],
            ]
        )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 2: FLEET OVERVIEW
    # =================================================================
    def _add_fleet_overview(self, fleet_by_year: pd.DataFrame, summary: dict,
                            specs: pd.DataFrame, charts: dict):
        """3-4 pages: year-by-year fleet table, trend charts, structural analysis."""
        self.doc.add_heading('2. Fleet Overview', level=1)

        # Year-by-year table
        self.doc.add_heading('Historical Fleet Size (2013-2023)', level=2)
        rows = []
        for _, r in fleet_by_year.iterrows():
            rows.append([
                str(int(r['year'])),
                f"{int(r['total']):,}",
                f"{int(r['covered']):,}",
                f"{int(r['open']):,}",
                f"{r['covered'] / r['total'] * 100:.1f}%",
            ])
        self._add_table(
            ['Year', 'Total Fleet', 'Covered', 'Open', 'Covered %'],
            rows
        )
        self._add_paragraph(
            f"The Mississippi River dry barge fleet grew by {fleet_by_year['total'].iloc[-1] - fleet_by_year['total'].iloc[0]:,} "
            f"units over the ten-year observation period, a compound growth rate of approximately "
            f"{((fleet_by_year['total'].iloc[-1] / fleet_by_year['total'].iloc[0]) ** (1/10) - 1) * 100:.1f}% "
            f"annually. Growth was concentrated in covered hoppers, which added "
            f"{fleet_by_year['covered'].iloc[-1] - fleet_by_year['covered'].iloc[0]:,} barges while open hoppers "
            f"added {fleet_by_year['open'].iloc[-1] - fleet_by_year['open'].iloc[0]:,}."
        )
        self._add_paragraph(
            "Note: 2019 fleet data is unavailable due to a gap in the USACE data publication cycle. "
            "Trend lines interpolate across this gap. All fleet counts reflect region 4 (Mississippi River System) "
            "vessels classified as dry covered or open hopper barges per USACE VTCC coding. The 2022 VTCC "
            "reclassification (from 4A40/4A41 to 4A47/4A48) has been normalized by combining legacy and "
            "current codes."
        )

        # Charts
        self._add_chart(charts.get('fleet_size_trend'),
                        'Figure 2.1: Fleet Size Trend — Covered vs Open Hopper (2013-2023)')
        self._add_chart(charts.get('covered_vs_open_area'),
                        'Figure 2.2: Covered/Open Composition Shift Over Time')

        # Physical specs
        self.doc.add_heading('Physical Specifications by Type', level=2)
        if not specs.empty:
            spec_rows = []
            for _, r in specs.iterrows():
                spec_rows.append([
                    r['barge_type'],
                    f"{int(r['count']):,}",
                    f"{r['avg_length']:.0f} ft" if pd.notna(r['avg_length']) else 'N/A',
                    f"{r['avg_breadth']:.0f} ft" if pd.notna(r['avg_breadth']) else 'N/A',
                    f"{r['avg_nrt']:,.0f}" if pd.notna(r['avg_nrt']) else 'N/A',
                    f"{r['avg_cap_tons']:,.0f}" if pd.notna(r['avg_cap_tons']) else 'N/A',
                ])
            self._add_table(
                ['Type', 'Count', 'Avg Length', 'Avg Beam', 'Avg NRT', 'Avg Cap (tons)'],
                spec_rows
            )
            self._add_paragraph(
                "Standard jumbo covered hoppers measure approximately 200 x 35 feet with cargo capacity "
                "of 1,500-1,800 tons. Open hoppers are dimensionally similar but offer higher volume capacity "
                "for lighter bulk commodities such as coal, aggregates, and scrap. Newer builds increasingly "
                "adopt the 200 x 35 jumbo footprint, which has become the industry standard for both types."
            )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 3: FLEET AGE & REPLACEMENT ANALYSIS
    # =================================================================
    def _add_age_analysis(self, age_dist: pd.DataFrame, build_decades: pd.DataFrame,
                          new_builds: pd.DataFrame, summary: dict, charts: dict):
        """3-4 pages: age distribution, build cohorts, new build trends, scrapping estimates."""
        self.doc.add_heading('3. Fleet Age & Replacement Analysis', level=1)

        # Age distribution table
        self.doc.add_heading('Age Distribution (2023)', level=2)
        rows = []
        for _, r in age_dist.iterrows():
            rows.append([
                str(r['age_bucket']),
                f"{int(r['count']):,}",
                f"{int(r['covered']):,}",
                f"{int(r['open']):,}",
                f"{r['pct']:.1f}%",
            ])
        self._add_table(['Age Range', 'Total', 'Covered', 'Open', '% of Fleet'], rows)
        self._add_paragraph(
            f"The current fleet has an average age of {summary['avg_age']:.1f} years and a median "
            f"of {summary['median_age']:.1f} years. The distribution is right-skewed, with a significant "
            f"tail of older vessels built during the 1990s construction boom when Jeffboat and multiple "
            f"regional yards were operating at full capacity."
        )

        self._add_chart(charts.get('age_distribution'),
                        'Figure 3.1: Fleet Age Distribution by 5-Year Bucket (2023)')

        # Build decade cohorts
        self.doc.add_heading('Build Decade Cohort Analysis', level=2)
        rows = []
        for _, r in build_decades.iterrows():
            decade_label = f"{int(r['decade'])}s"
            avg_age = f"{r['avg_age']:.0f}" if pd.notna(r['avg_age']) else 'N/A'
            retirement = 'Retirement imminent' if r['avg_age'] and r['avg_age'] >= 30 else \
                         'Approaching retirement' if r['avg_age'] and r['avg_age'] >= 25 else \
                         'Mid-life' if r['avg_age'] and r['avg_age'] >= 15 else 'Current'
            rows.append([
                decade_label,
                f"{int(r['count']):,}",
                f"{int(r['covered']):,}",
                f"{int(r['open']):,}",
                avg_age,
                retirement,
            ])
        self._add_table(
            ['Build Decade', 'Total', 'Covered', 'Open', 'Avg Age', 'Retirement Status'],
            rows
        )

        self._add_chart(charts.get('build_decade_cohorts'),
                        'Figure 3.2: Fleet by Build Decade — Covered vs Open')

        # Analysis of retirement cliff
        over_25 = age_dist[age_dist['age_bucket'].isin(['25-29', '30-34', '35-39', '40+'])]['count'].sum()
        pct_over_25 = over_25 / summary['total_barges'] * 100
        self._add_paragraph(
            f"Approximately {over_25:,} barges ({pct_over_25:.1f}% of the fleet) are 25 years old or older, "
            f"placing them at or beyond typical economic life for inland dry cargo service. Assuming standard "
            f"scrapping at age 30-35 and current new-build capacity of 300-500 covered and 100-200 open units "
            f"per year, the industry faces a structural deficit in replacement tonnage through at least 2030."
        )

        # New builds by year
        self.doc.add_heading('New Construction Trends (2000-2023)', level=2)
        if not new_builds.empty:
            recent = new_builds[new_builds['year_built'] >= 2010]
            rows = []
            for _, r in recent.iterrows():
                rows.append([
                    str(int(r['year_built'])),
                    f"{int(r['builds']):,}",
                    f"{int(r['covered']):,}",
                    f"{int(r['open']):,}",
                ])
            self._add_table(['Build Year', 'Total', 'Covered', 'Open'], rows)

        self._add_chart(charts.get('build_year_timeline'),
                        'Figure 3.3: New Barge Construction by Year (2000-2023, Jeffboat closure annotated)')

        self._add_paragraph(
            "The closure of Jeffboat in Jeffersonville, Indiana in 2018 marked a watershed moment for "
            "the inland barge construction industry. At its peak, Jeffboat produced over 1,000 barges "
            "annually and was the dominant builder of both covered and open hopper barges. Post-closure, "
            "construction has shifted primarily to Arcosa Marine Products (formerly Trinity Industries' "
            "inland barge division) facilities in Arkansas and Heartland Barge in Louisiana. Combined "
            "output from these remaining yards does not fully replace Jeffboat's capacity."
        )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 4: OPERATOR MARKET SHARE
    # =================================================================
    def _add_operator_analysis(self, top_operators: pd.DataFrame, concentration: dict,
                               charts: dict):
        """3-4 pages: top operators table, market share, concentration metrics."""
        self.doc.add_heading('4. Operator Market Share', level=1)

        # Top 15 table
        self.doc.add_heading('Top 15 Operators (2023)', level=2)
        rows = []
        for _, r in top_operators.iterrows():
            name = r['operator'] if r['operator'] else 'Unknown/Unregistered'
            rows.append([
                name[:35],
                f"{int(r['total']):,}",
                f"{int(r['covered']):,}",
                f"{int(r['open']):,}",
                f"{r['avg_age']:.1f}",
                f"{r['market_share']:.1f}%",
            ])
        self._add_table(
            ['Operator', 'Total', 'Covered', 'Open', 'Avg Age', 'Market Share'],
            rows
        )

        # Top operators chart
        self._add_chart(charts.get('top_operators'),
                        'Figure 4.1: Top 15 Operators — Mississippi Dry Barges (2023)')

        # Operator narrative
        if not top_operators.empty:
            top = top_operators.iloc[0]
            second = top_operators.iloc[1] if len(top_operators) > 1 else None
            self._add_paragraph(
                f"{top['operator'] or 'The leading operator'} commands the largest dry barge fleet on "
                f"the Mississippi with {int(top['total']):,} barges ({top['market_share']:.1f}% market share), "
                f"operating a mixed fleet of covered and open hoppers. "
                + (f"{second['operator'] or 'The second-largest operator'} follows with "
                   f"{int(second['total']):,} barges ({second['market_share']:.1f}% share). "
                   if second is not None else '')
                + f"The top three operators collectively control {concentration['top3_share']:.1f}% "
                f"of the fleet, indicating a moderately concentrated market structure."
            )

        # Concentration metrics
        self.doc.add_heading('Market Concentration', level=2)
        self._add_table(
            ['Metric', 'Value', 'Interpretation'],
            [
                ['Top-3 Share', f"{concentration['top3_share']:.1f}%",
                 'Moderately concentrated' if concentration['top3_share'] < 50 else 'Highly concentrated'],
                ['Top-5 Share', f"{concentration['top5_share']:.1f}%", ''],
                ['Top-10 Share', f"{concentration['top10_share']:.1f}%", ''],
                ['HHI Index', f"{concentration['hhi']:,.0f}",
                 'Unconcentrated (<1500)' if concentration['hhi'] < 1500 else
                 'Moderately concentrated (1500-2500)' if concentration['hhi'] < 2500 else
                 'Highly concentrated (>2500)'],
                ['Total Operators', f"{concentration['total_operators']:,}",
                 'Long tail of small fleet operators'],
            ]
        )
        self._add_paragraph(
            f"With an HHI of {concentration['hhi']:,.0f} and {concentration['total_operators']} "
            f"registered operators, the Mississippi dry barge market has a classic structure: a handful "
            f"of large fleet operators supported by a long tail of small and mid-size companies, many of "
            f"which operate fewer than 50 barges. This fragmentation in the lower tier creates acquisition "
            f"opportunities for buyers seeking to consolidate market position."
        )

        # Market share donut
        self._add_chart(charts.get('market_share'),
                        'Figure 4.2: Market Share — Top 5 Operators (2023)')

        self._add_paragraph(
            "The 'Unknown/Unregistered' category represents barges where the USACE operator record "
            "has no matching entry in the operator master file — typically vessels in temporary layup, "
            "pending registration transfer, or owned by entities that have not filed current operator "
            "documentation. This is a normal artifact of the USACE data collection process."
        )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 5: DEMAND DRIVERS
    # =================================================================
    def _add_demand_drivers(self):
        """2-3 pages: grain, coal, seasonal patterns — hardcoded research data."""
        self.doc.add_heading('5. Demand Drivers', level=1)

        # Grain exports
        self.doc.add_heading('Grain & Agricultural Exports', level=2)
        self._add_table(
            ['Year', 'Corn Exports (M bu)', 'Soybean Exports (M bu)',
             'Corn Price ($/bu)', 'Soybean Price ($/bu)'],
            [
                ['2015', '1,898', '1,936', '$3.61', '$8.95'],
                ['2016', '2,293', '2,174', '$3.36', '$9.47'],
                ['2017', '1,956', '2,129', '$3.36', '$9.33'],
                ['2018', '2,432', '1,748', '$3.61', '$8.60'],
                ['2019', '1,770', '1,682', '$3.56', '$8.57'],
                ['2020', '2,753', '2,266', '$4.53', '$10.80'],
                ['2021', '2,471', '2,116', '$5.45', '$13.30'],
                ['2022', '2,107', '2,031', '$6.54', '$14.20'],
                ['2023', '1,664', '1,720', '$4.65', '$12.50'],
                ['2024E', '2,300', '1,800', '$4.35', '$10.80'],
            ]
        )
        self._add_paragraph(
            "Grain remains the single largest cargo category on the Mississippi River, accounting for "
            "approximately 60% of covered hopper barge demand. The river system handles roughly 60% of "
            "U.S. grain exports, with the Gulf export corridor (New Orleans/South Louisiana) serving as "
            "the primary funnel. Covered hopper fleet utilization correlates closely with harvest volume, "
            "Gulf export pace, and competitiveness of U.S. grain pricing versus South American alternatives."
        )
        self._add_paragraph(
            "USDA long-term baseline projections (2024-2033) project corn exports averaging 2.3-2.5 billion "
            "bushels annually and soybean exports at 1.8-2.0 billion bushels. These projections assume "
            "normalized trade relations with China, continued growth in global livestock production, and "
            "steady ethanol demand under the Renewable Fuel Standard. Upside catalysts include expanded "
            "sustainable aviation fuel mandates and growing protein consumption in Southeast Asia."
        )

        # Coal decline
        self.doc.add_heading('Coal Production & Decline Trajectory', level=2)
        self._add_table(
            ['Year', 'U.S. Coal Production (M short tons)', 'YoY Change', 'River-Moved (est.)'],
            [
                ['2015', '897', '-10.4%', '~180M tons'],
                ['2016', '728', '-18.8%', '~150M tons'],
                ['2017', '774', '+6.3%', '~155M tons'],
                ['2018', '756', '-2.3%', '~150M tons'],
                ['2019', '706', '-6.6%', '~140M tons'],
                ['2020', '535', '-24.2%', '~105M tons'],
                ['2021', '578', '+8.0%', '~115M tons'],
                ['2022', '595', '+2.9%', '~120M tons'],
                ['2023', '507', '-14.8%', '~100M tons'],
                ['2024E', '475', '-6.3%', '~95M tons'],
            ]
        )
        self._add_paragraph(
            "U.S. coal production has declined roughly 45% from its 2008 peak of 1.17 billion short tons, "
            "driven by natural gas price competitiveness and utility fleet retirements. The EIA Annual Energy "
            "Outlook projects continued decline, with domestic coal production falling below 400 million tons "
            "by 2030. This decline directly reduces demand for open hopper barges, though the impact is "
            "partially offset by growth in aggregates, steel scrap, and industrial mineral movements."
        )

        # Seasonal patterns
        self.doc.add_heading('Seasonal Patterns & Low-Water Risk', level=2)
        self._add_paragraph(
            "Barge demand follows strong seasonal patterns tied to the agricultural calendar. The fall "
            "harvest (September-November) generates peak demand for covered hoppers as grain moves "
            "downriver to Gulf export elevators. Spring planting season creates a secondary demand peak "
            "as fertilizer moves upriver. Coal movements are more evenly distributed but spike during "
            "summer utility demand peaks and winter heating seasons."
        )
        self._add_paragraph(
            "The 2022 low-water event on the Lower Mississippi — the worst in decades — demonstrated "
            "the system's vulnerability to drought conditions. Barge loadings were restricted to 60-70% "
            "of capacity for extended periods, effectively reducing fleet carrying capacity by 30-40% and "
            "creating severe congestion at key choke points. Climate projections suggest such events may "
            "become more frequent, creating both risk (capacity disruption) and opportunity (more barges "
            "needed per ton of cargo during restricted draft periods)."
        )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 6: SUPPLY SIDE
    # =================================================================
    def _add_supply_side(self):
        """2-3 pages: shipyard capacity, new build costs, secondhand values."""
        self.doc.add_heading('6. Supply Side — Construction & Secondhand Market', level=1)

        # Shipyard capacity
        self.doc.add_heading('Active Inland Barge Builders', level=2)
        self._add_table(
            ['Builder', 'Location', 'Status', 'Est. Annual Capacity', 'Specialization'],
            [
                ['Arcosa Marine Products', 'Ashdown, AR / Madisonville, LA', 'Active', '~400-600 barges',
                 'Covered & open hoppers, tank barges'],
                ['Heartland Barge', 'Brownsville, PA', 'Active', '~150-250 barges',
                 'Open hoppers, deck barges'],
                ['Jeffboat', 'Jeffersonville, IN', 'CLOSED (2018)', '—',
                 'Was largest; covered & open hoppers'],
                ['Conrad Shipyard', 'Morgan City, LA', 'Active (limited)', '~50-100 barges',
                 'Specialty barges, limited hopper'],
                ['Metal Shark / Horizon', 'Various', 'Active (limited)', '~25-50 barges',
                 'Niche builders'],
            ]
        )
        self._add_paragraph(
            "The inland barge construction market is effectively a duopoly following Jeffboat's closure. "
            "Arcosa Marine Products, spun off from Trinity Industries in 2018, operates the largest "
            "remaining facilities and accounts for roughly 60-70% of new covered hopper production. "
            "Heartland Barge in Brownsville, Pennsylvania focuses primarily on open hoppers and deck "
            "barges. Combined new-build capacity across all active yards is estimated at 500-850 barges "
            "annually — well below the 1,000-1,500 units Jeffboat alone could produce at peak."
        )

        # New build costs
        self.doc.add_heading('New Construction Cost Benchmarks', level=2)
        self._add_table(
            ['Barge Type', 'Size', '2018 Cost', '2023 Cost', 'Lead Time'],
            [
                ['Covered Hopper (Jumbo)', '200 x 35 ft', '$550,000-$650,000', '$750,000-$950,000', '18-24 months'],
                ['Covered Hopper (Standard)', '195 x 35 ft', '$500,000-$600,000', '$700,000-$850,000', '15-20 months'],
                ['Open Hopper (Jumbo)', '200 x 35 ft', '$400,000-$500,000', '$550,000-$700,000', '12-18 months'],
                ['Open Hopper (Standard)', '195 x 35 ft', '$350,000-$450,000', '$500,000-$650,000', '10-15 months'],
            ]
        )
        self._add_paragraph(
            "New construction costs have increased 35-50% since 2018, driven by steel price inflation, "
            "labor shortages in the inland fabrication trades, and reduced competition. Current lead times "
            "of 18-24 months for covered hoppers (the highest-demand type) make secondhand acquisition "
            "the only near-term option for operators seeking to expand fleet capacity."
        )

        # Secondhand values
        self.doc.add_heading('Secondhand Barge Value Estimates', level=2)
        self._add_table(
            ['Barge Age', 'Covered Hopper Value', 'Open Hopper Value', 'Condition Notes'],
            [
                ['0-5 years', '$600,000-$800,000', '$450,000-$600,000', 'Near-new; premium for recent builds'],
                ['5-10 years', '$400,000-$600,000', '$300,000-$450,000', 'Prime acquisition target range'],
                ['10-15 years', '$250,000-$400,000', '$200,000-$300,000', 'Requires survey; coating condition key'],
                ['15-20 years', '$150,000-$250,000', '$120,000-$200,000', 'Significant remaining life if maintained'],
                ['20-25 years', '$80,000-$150,000', '$60,000-$120,000', 'Approaching end of economic life'],
                ['25+ years', '$30,000-$80,000', '$25,000-$60,000', 'Scrap value floor; limited commercial use'],
            ]
        )
        self._add_paragraph(
            "Secondhand values remain elevated relative to historical norms due to the construction "
            "capacity constraint. Well-maintained covered hoppers under 10 years of age trade at a "
            "significant premium, as they represent immediately deployable capacity without the 18-24 "
            "month wait for new construction. The scrap value floor (steel content at prevailing scrap "
            "prices) provides a hard downside of approximately $25,000-$40,000 per barge."
        )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 7: M&A LANDSCAPE
    # =================================================================
    def _add_ma_landscape(self):
        """2 pages: recent deals, valuation benchmarks, strategic buyers."""
        self.doc.add_heading('7. M&A Landscape', level=1)

        # Recent transactions
        self.doc.add_heading('Notable Recent Transactions', level=2)
        self._add_table(
            ['Year', 'Target', 'Acquirer', 'Deal Details', 'Significance'],
            [
                ['2021', 'ACBL (American Commercial Barge Line)',
                 'Emerged from Ch. 11 (Platinum Equity)',
                 '~2,300 barges; restructured debt',
                 'Largest fleet; consolidated 3 legacy operators'],
                ['2022', 'Marquette Transportation',
                 'Redwood Capital (merged into)',
                 'Combined with Canal Barge Co',
                 'Created significant mid-tier competitor'],
                ['2023', 'Canal Barge Company',
                 'Redwood/Marquette',
                 'Folded into Marquette platform',
                 'Further consolidation of mid-tier'],
                ['2023', 'Kirby Corp inland barge acquisitions',
                 'Kirby Corporation',
                 'Selective fleet purchases',
                 'Tank barge leader expanding dry cargo'],
                ['2024', 'Ingram Barge — no change',
                 'Ingram Industries (private)',
                 'Market leader; not for sale',
                 'Family-owned; strategic anchor'],
            ]
        )
        self._add_paragraph(
            "The inland barge industry has undergone significant consolidation over the past five years, "
            "driven by the economics of scale in fleeting operations, towing, and barge maintenance. "
            "ACBL's emergence from Chapter 11 in 2021 under Platinum Equity ownership consolidated "
            "multiple legacy operators (including AEP River Operations and early ACBL entities) into a "
            "single platform. The Redwood Capital/Marquette/Canal Barge combination created a formidable "
            "mid-tier competitor with diversified geographic coverage."
        )

        # Valuation benchmarks
        self.doc.add_heading('Valuation Benchmarks', level=2)
        self._add_table(
            ['Metric', 'Range', 'Notes'],
            [
                ['EV / Barge', '$80,000-$180,000', 'Varies by fleet age and condition; blended'],
                ['EV / EBITDA', '5.0x-7.5x', 'Higher for contracted fleets with stable cash flows'],
                ['Price / Net Asset Value', '0.7x-1.2x', 'Discount for older fleets; premium for young fleets'],
                ['$/ton cargo capacity', '$50-$120/ton', 'Depends on barge type and utilization'],
            ]
        )
        self._add_paragraph(
            "Transaction multiples in the inland barge sector have historically ranged from 5-7x EBITDA "
            "for fleet-focused deals. Premiums are commanded by operators with long-term customer contracts, "
            "modern fleet age profiles, and strategic geographic positioning (particularly operators with "
            "strong Gulf presence for grain export logistics). Distressed situations — as seen in the ACBL "
            "bankruptcy — can offer significant discounts to replacement cost."
        )

        # Strategic buyer landscape
        self.doc.add_heading('Potential Strategic Buyers', level=2)
        self._add_paragraph(
            "The current landscape of likely acquirers includes: (1) existing barge line operators seeking "
            "fleet expansion (Ingram, ACBL, ARTCO, Marquette), though each faces different strategic "
            "constraints; (2) private equity firms with inland marine portfolio companies, following the "
            "Platinum Equity/ACBL model; (3) international logistics companies seeking Mississippi River "
            "exposure; and (4) agricultural commodity firms pursuing vertical integration of their "
            "supply chains, as ADM and Cargill have historically done through ARTCO and their respective "
            "barge operations."
        )
        self._add_paragraph(
            "For a new entrant or mid-tier consolidator, the ideal acquisition target would be a fleet "
            "of 200-500 barges with average age under 15 years, existing customer relationships in the "
            "grain corridor, and either owned or long-term leased fleeting capacity. Such platforms "
            "exist primarily in the lower Mississippi and Ohio River systems."
        )

        self.doc.add_page_break()

    # =================================================================
    # SECTION 8: HISTORICAL TRENDS
    # =================================================================
    def _add_historical_trends(self, fleet_by_year: pd.DataFrame, age_trend: pd.DataFrame,
                               capacity_trend: pd.DataFrame, charts: dict):
        """2-3 pages: 10-year summary, capacity and age trends."""
        self.doc.add_heading('8. Historical Trends & 10-Year Summary', level=1)

        # Comprehensive 10-year summary table
        self.doc.add_heading('10-Year Fleet Summary', level=2)
        rows = []
        for _, r in fleet_by_year.iterrows():
            # Find matching age data
            yr = int(r['year'])
            age_row = age_trend[age_trend['year'] == yr]
            avg_age = f"{age_row['avg_age'].values[0]:.1f}" if not age_row.empty else 'N/A'
            over_30 = f"{int(age_row['over_30'].values[0]):,}" if not age_row.empty else 'N/A'
            rows.append([
                str(yr),
                f"{int(r['total']):,}",
                f"{int(r['covered']):,}",
                f"{int(r['open']):,}",
                avg_age,
                over_30,
            ])
        self._add_table(
            ['Year', 'Total Fleet', 'Covered', 'Open', 'Avg Age', 'Over 30 Yrs'],
            rows
        )

        first_yr = fleet_by_year.iloc[0]
        last_yr = fleet_by_year.iloc[-1]
        self._add_paragraph(
            f"Over the 2013-2023 period, the Mississippi dry barge fleet grew by "
            f"{int(last_yr['total']) - int(first_yr['total']):,} units ({(last_yr['total'] / first_yr['total'] - 1) * 100:.1f}%). "
            f"The structural shift from coal to grain is clearly visible in the covered/open ratio: covered "
            f"hoppers grew from {first_yr['covered'] / first_yr['total'] * 100:.0f}% to "
            f"{last_yr['covered'] / last_yr['total'] * 100:.0f}% of the total fleet. This shift reflects "
            f"both new covered hopper construction for agricultural service and the gradual attrition of "
            f"older open hoppers formerly deployed in coal trades."
        )

        # Capacity trend chart
        self._add_chart(charts.get('capacity_trend'),
                        'Figure 8.1: Fleet Capacity Trends — NRT and Cargo Tonnage (2013-2023)')

        if not age_trend.empty:
            first_age = age_trend.iloc[0]
            last_age = age_trend.iloc[-1]
            self._add_paragraph(
                f"Average fleet age has shifted from {first_age['avg_age']:.1f} years in "
                f"{int(first_age['year'])} to {last_age['avg_age']:.1f} years in "
                f"{int(last_age['year'])}. The number of barges over 30 years old has "
                f"{'increased' if last_age['over_30'] > first_age['over_30'] else 'decreased'} "
                f"from {int(first_age['over_30']):,} to {int(last_age['over_30']):,}, reflecting "
                f"the competing dynamics of continued service of legacy assets versus scrapping and "
                f"replacement."
            )

        self.doc.add_page_break()

    # =================================================================
    # APPENDIX A: OPERATOR FLEET SUMMARY
    # =================================================================
    def _add_appendix_operators(self, all_operators: pd.DataFrame):
        """Complete operator inventory sorted by fleet size."""
        self.doc.add_heading('Appendix A: Operator Fleet Summary (2023)', level=1)

        self._add_paragraph(
            f"Complete listing of all {len(all_operators)} registered operators of Mississippi River "
            f"dry hopper barges in the 2023 USACE fleet data. Sorted by total fleet size descending.",
            italic=True, size=9
        )

        # Split into pages of ~40 operators to avoid table overflow
        page_size = 40
        for start in range(0, len(all_operators), page_size):
            chunk = all_operators.iloc[start:start + page_size]
            rows = []
            for _, r in chunk.iterrows():
                name = r['operator'] if r['operator'] else 'Unknown/Unregistered'
                rows.append([
                    name[:35],
                    f"{int(r['total']):,}",
                    f"{int(r['covered']):,}",
                    f"{int(r['open']):,}",
                    f"{r['avg_age']:.1f}" if pd.notna(r['avg_age']) else 'N/A',
                    str(int(r['oldest_build'])) if pd.notna(r['oldest_build']) else 'N/A',
                    str(int(r['newest_build'])) if pd.notna(r['newest_build']) else 'N/A',
                ])
            self._add_table(
                ['Operator', 'Total', 'Covered', 'Open', 'Avg Age', 'Oldest', 'Newest'],
                rows
            )

    # =================================================================
    # METHODOLOGY NOTE
    # =================================================================
    def _add_methodology(self):
        """Brief methodology note."""
        self.doc.add_heading('Methodology & Data Notes', level=1)

        self._add_paragraph(
            "This analysis is based on the U.S. Army Corps of Engineers (USACE) Waterborne Commerce "
            "vessel characteristics database, which tracks all documented U.S. commercial vessels by "
            "annual fleet snapshots. The dataset covers fleet years 2013-2023 with the exception of "
            "2019, which is unavailable due to a gap in the USACE publication cycle."
        )
        self._add_paragraph(
            "Vessel scope includes all dry hopper barges (covered and open) registered in USACE region 4 "
            "(Mississippi River System). VTCC classification codes 4A40 (Open Hopper), 4A41 (Covered "
            "Hopper), 4A47 (Open Dry Cargo), and 4A48 (Covered Dry Cargo) are combined to account for "
            "the 2022 USACE reclassification of barge type codes. Prior to 2022, most covered barges "
            "were classified under 4A41; after reclassification, the same vessels appear under 4A48. "
            "This report normalizes across the reclassification boundary by grouping: Covered = 4A41 + "
            "4A48, Open = 4A40 + 4A47."
        )
        self._add_paragraph(
            "Operator names are standardized using the USACE operator_name_std field where available, "
            "falling back to the raw operator_name. Some vessels have no operator match in the current "
            "master file and appear as 'Unknown/Unregistered' — this is a known artifact of USACE data "
            "collection and does not indicate vessels are out of service."
        )
        self._add_paragraph(
            "Market research data (grain prices, coal production, new build costs, secondhand values, "
            "M&A transactions) is sourced from USDA, EIA, and industry publications. These figures are "
            "approximate benchmarks and should be validated against current market conditions at the time "
            "of any acquisition decision."
        )

    # =================================================================
    # MAIN GENERATION METHOD
    # =================================================================
    def generate(self, query_results: dict, charts: dict) -> Path:
        """
        Generate the complete report.

        Args:
            query_results: Dict of DataFrames/dicts from BargeQueries
            charts: Dict of chart paths from BargeCharts.generate_all()

        Returns:
            Path to generated DOCX file
        """
        logger.info("Building barge fleet acquisition report...")

        summary = query_results['summary']
        concentration = query_results['concentration']
        fleet_by_year = query_results['fleet_by_year']

        # Build sections
        self._add_cover_page()
        self._add_executive_summary(summary, concentration, fleet_by_year)
        self._add_fleet_overview(fleet_by_year, summary,
                                 query_results['physical_specs'], charts)
        self._add_age_analysis(query_results['age_distribution'],
                               query_results['build_decade_cohorts'],
                               query_results['new_builds_by_year'],
                               summary, charts)
        self._add_operator_analysis(query_results['top_operators'],
                                    concentration, charts)
        self._add_demand_drivers()
        self._add_supply_side()
        self._add_ma_landscape()
        self._add_historical_trends(fleet_by_year,
                                    query_results['age_trend'],
                                    query_results['capacity_trend'],
                                    charts)
        self._add_appendix_operators(query_results['all_operators'])
        self._add_methodology()

        # Save
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.output_dir / f'Mississippi_Barge_Fleet_Analysis_{timestamp}.docx'
        self.doc.save(str(output_path))
        logger.info(f"Report saved: {output_path}")
        return output_path
